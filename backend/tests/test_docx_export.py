"""DOCX export — build_docx unit + /export/docx endpoint (wizard & freeform)."""

import base64
import io

import pytest
from docx import Document

from app.services.docx_export import build_docx

DOCX_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
MD = "# Бриф\n\n## Цель\n\n- первый **пункт**\n\nфинальный абзац"

# 1x1 PNG, который распознаёт python-docx (для logo-тестов; без интернета)
PNG_BYTES = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAIAAACQd1PeAAAADElEQVR4nGP4z8AAAAMBAQDJ"
    "/pLvAAAAAElFTkSuQmCC"
)


# --- unit (no DB) ---------------------------------------------------------


def test_build_docx_returns_valid_docx():
    data = build_docx(MD, title="Тестовый бриф")
    assert data[:4] == b"PK\x03\x04"  # zip/docx signature

    doc = Document(io.BytesIO(data))
    styles = [p.style.name for p in doc.paragraphs]
    assert "Heading 1" in styles
    assert "Heading 2" in styles
    assert "List Bullet" in styles
    # инлайн **bold** → жирный run
    assert any(run.bold for p in doc.paragraphs for run in p.runs)
    # title попал в core-properties, но не задублирован в тело
    assert doc.core_properties.title == "Тестовый бриф"


def test_build_docx_empty_is_still_valid():
    data = build_docx("")
    assert data[:4] == b"PK\x03\x04"
    Document(io.BytesIO(data))  # открывается без ошибок


def _first_heading_color(doc):
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading") and p.runs:
            return p.runs[0].font.color.rgb
    return None


def test_build_docx_with_identity_applies_font_and_color():
    identity = {"font_family": "Georgia", "accent_color": "#1f6feb"}
    doc = Document(io.BytesIO(build_docx(MD, title="t", identity=identity)))
    assert doc.styles["Normal"].font.name == "Georgia"
    # цвет акцента ушёл в заголовки
    assert str(_first_heading_color(doc)) == "1F6FEB"


def test_build_docx_identity_color_falls_back_to_primary():
    identity = {"primary_color": "#abc"}  # #RGB → #RRGGBB, accent отсутствует
    doc = Document(io.BytesIO(build_docx(MD, identity=identity)))
    assert str(_first_heading_color(doc)) == "AABBCC"


@pytest.mark.parametrize(
    "identity",
    [
        None,
        {},
        {"font_family": "", "primary_color": "", "accent_color": None},
        {"primary_color": "not-a-color", "accent_color": "#xyz"},
    ],
)
def test_build_docx_empty_or_invalid_identity_does_not_break(identity):
    data = build_docx(MD, title="t", identity=identity)
    assert data[:4] == b"PK\x03\x04"
    doc = Document(io.BytesIO(data))
    # битый/пустой цвет → заголовки без явного цвета (как раньше)
    assert _first_heading_color(doc) is None


def _all_text(doc):
    return "\n".join(p.text for p in doc.paragraphs)


def test_build_docx_header_block_renders():
    doc = Document(
        io.BytesIO(
            build_docx(
                MD,
                title="t",
                brand_name="Север",
                document_label="Креативный бриф",
                date_text="25.06.2026",
                identity={"font_family": "Georgia", "accent_color": "#1f6feb"},
            )
        )
    )
    text = _all_text(doc)
    assert "Север" in text
    assert "Креативный бриф" in text
    assert "25.06.2026" in text
    # бренд-шрифт всё ещё применяется к телу
    assert doc.styles["Normal"].font.name == "Georgia"


def test_build_docx_without_header_kwargs_has_no_letterhead():
    doc = Document(io.BytesIO(build_docx(MD, title="t")))
    # первый непустой параграф — markdown H1, а не letterhead
    first = next(p for p in doc.paragraphs if p.text.strip())
    assert first.style.name == "Heading 1"
    assert first.text.strip() == "Бриф"
    # нет meta-разделителя
    assert " · " not in _all_text(doc)


def test_build_docx_with_logo_embeds_image():
    doc = Document(
        io.BytesIO(build_docx(MD, brand_name="Север", logo_bytes=PNG_BYTES))
    )
    assert len(doc.inline_shapes) == 1
    # бренд и тело документа на месте
    text = _all_text(doc)
    assert "Север" in text
    assert "Бриф" in text


def test_build_docx_without_logo_has_no_image():
    doc = Document(io.BytesIO(build_docx(MD, brand_name="Север", logo_bytes=None)))
    assert len(doc.inline_shapes) == 0


def test_build_docx_broken_logo_does_not_break():
    doc = Document(
        io.BytesIO(
            build_docx(MD, brand_name="Север", logo_bytes=b"definitely not an image")
        )
    )
    assert len(doc.inline_shapes) == 0  # битый логотип пропущен
    text = _all_text(doc)
    assert "Север" in text  # header/body не сломаны
    assert "Бриф" in text


# --- endpoint (DB) --------------------------------------------------------


def _generated_wizard(db_session):
    from app.models import Brief

    brief = Brief(title="W docx", generated_markdown=MD, status="generated")
    db_session.add(brief)
    db_session.commit()
    return brief.id


def _generated_freeform(db_session, identity=None):
    from app.models import Brand, Brief

    brand = Brand(name="B", brand_identity_json=identity or {})
    db_session.add(brand)
    db_session.flush()
    brief = Brief(
        title="F docx", brand_id=brand.id, generated_markdown=MD, status="generated"
    )
    db_session.add(brief)
    db_session.commit()
    return brief.id


@pytest.mark.db
def test_export_docx_409_when_not_generated(client):
    brief_id = client.post("/api/briefs", json={"title": "x"}).json()["id"]
    assert client.get(f"/api/briefs/{brief_id}/export/docx").status_code == 409


@pytest.mark.db
def test_export_docx_wizard(client, db_session):
    brief_id = _generated_wizard(db_session)
    resp = client.get(f"/api/briefs/{brief_id}/export/docx")
    assert resp.status_code == 200
    assert resp.headers["content-type"] == DOCX_CT
    assert resp.headers["content-disposition"] == (
        f'attachment; filename="brief-{brief_id}.docx"'
    )
    assert resp.content[:4] == b"PK\x03\x04"
    Document(io.BytesIO(resp.content))  # валидный docx


@pytest.mark.db
def test_export_docx_freeform(client, db_session):
    brief_id = _generated_freeform(db_session)
    resp = client.get(f"/api/briefs/{brief_id}/export/docx")
    assert resp.status_code == 200
    assert resp.headers["content-type"] == DOCX_CT
    assert resp.content[:4] == b"PK\x03\x04"


@pytest.mark.db
def test_export_docx_freeform_with_identity(client, db_session):
    brief_id = _generated_freeform(
        db_session, identity={"font_family": "Georgia", "accent_color": "#1f6feb"}
    )
    resp = client.get(f"/api/briefs/{brief_id}/export/docx")
    assert resp.status_code == 200
    assert resp.content[:4] == b"PK\x03\x04"
    doc = Document(io.BytesIO(resp.content))
    assert doc.styles["Normal"].font.name == "Georgia"


@pytest.mark.db
def test_export_docx_header_has_brand_and_label(client, db_session):
    from app.models import Brand, Brief

    brand = Brand(name="Бренд-Икс", brand_identity_json={"accent_color": "#1f6feb"})
    db_session.add(brand)
    db_session.flush()
    brief = Brief(
        title="F", brand_id=brand.id, generated_markdown=MD, status="generated"
    )
    db_session.add(brief)
    db_session.commit()
    resp = client.get(f"/api/briefs/{brief.id}/export/docx")
    assert resp.status_code == 200
    text = "\n".join(p.text for p in Document(io.BytesIO(resp.content)).paragraphs)
    assert "Бренд-Икс" in text  # бренд-строка letterhead
    assert "Коммуникационный бриф" in text  # document_label (brief_type=custom)


@pytest.mark.db
def test_export_docx_no_brand_has_label_but_no_brand_name(client, db_session):
    brief_id = _generated_wizard(db_session)  # без бренда
    resp = client.get(f"/api/briefs/{brief_id}/export/docx")
    assert resp.status_code == 200
    doc = Document(io.BytesIO(resp.content))
    text = "\n".join(p.text for p in doc.paragraphs)
    # letterhead с типом документа есть и без бренда
    assert "Коммуникационный бриф" in text


@pytest.mark.db
def test_export_docx_embeds_logo_when_fetch_succeeds(client, db_session, monkeypatch):
    monkeypatch.setattr("app.routers.briefs.fetch_logo_bytes", lambda url: PNG_BYTES)
    brief_id = _generated_freeform(
        db_session, identity={"logo_url": "https://cdn.example.com/logo.png"}
    )
    resp = client.get(f"/api/briefs/{brief_id}/export/docx")
    assert resp.status_code == 200
    doc = Document(io.BytesIO(resp.content))
    assert len(doc.inline_shapes) == 1


@pytest.mark.db
def test_export_docx_no_logo_when_fetch_returns_none(client, db_session, monkeypatch):
    monkeypatch.setattr("app.routers.briefs.fetch_logo_bytes", lambda url: None)
    brief_id = _generated_freeform(
        db_session, identity={"logo_url": "https://cdn.example.com/logo.png"}
    )
    resp = client.get(f"/api/briefs/{brief_id}/export/docx")
    assert resp.status_code == 200
    doc = Document(io.BytesIO(resp.content))
    assert len(doc.inline_shapes) == 0
    # letterhead (бренд/тип) всё равно есть
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "B" in text  # brand name из _generated_freeform
    assert "Коммуникационный бриф" in text


@pytest.mark.db
def test_markdown_and_json_export_regression(client, db_session):
    brief_id = _generated_wizard(db_session)
    md = client.get(f"/api/briefs/{brief_id}/export/markdown")
    assert md.status_code == 200 and md.headers["content-type"].startswith("text/markdown")
    js = client.get(f"/api/briefs/{brief_id}/export/json")
    assert js.status_code == 200 and js.headers["content-type"].startswith("application/json")
