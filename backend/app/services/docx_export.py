"""Лёгкий конвертер Markdown → DOCX для экспорта сгенерированного брифа.

Покрывает подмножество markdown, которое выдают generate-промпты и рендерит
фронтовый MarkdownView: `#`/`##`/`###` заголовки, списки `- `/`* `, инлайн
`**bold**`, обычные параграфы. Таблицы не парсятся (генератор их не выдаёт).
Чистый модуль — без БД/HTTP.
"""

import io
import logging
import re
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

logger = logging.getLogger(__name__)

_HEADING_RE = re.compile(r"^(#{1,3})\s+(.*)$")
_BULLET_RE = re.compile(r"^[-*]\s+(.*)$")
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_HEX_RE = re.compile(r"^#(?:[0-9a-fA-F]{3}|[0-9a-fA-F]{6})$")

_META_GREY = RGBColor(0x88, 0x88, 0x88)
_NEUTRAL_DIVIDER = RGBColor(0xCC, 0xCC, 0xCC)


def _parse_hex(value: Any) -> RGBColor | None:
    """`#RGB` / `#RRGGBB` → RGBColor; всё прочее (None/пусто/мусор) → None."""
    if not isinstance(value, str):
        return None
    v = value.strip()
    if not _HEX_RE.match(v):
        return None
    h = v[1:]
    if len(h) == 3:  # #RGB → #RRGGBB
        h = "".join(c * 2 for c in h)
    try:
        return RGBColor.from_string(h.upper())
    except ValueError:
        return None


def _add_divider(document, color: RGBColor) -> None:
    """Тонкая горизонтальная линия — нижняя граница пустого параграфа."""
    paragraph = document.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")  # 1/8 pt → 12 = 1.5pt
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), str(color))  # RGBColor → 'RRGGBB'
    borders.append(bottom)
    p_pr.append(borders)


def _add_logo(document, logo_bytes: bytes) -> None:
    """Вставить логотип отдельным параграфом сверху.

    Если python-docx не распознал картинку (или иная ошибка) — пустой параграф
    удаляется, экспорт продолжается без логотипа.
    """
    paragraph = document.add_paragraph()
    try:
        paragraph.add_run().add_picture(io.BytesIO(logo_bytes), width=Inches(1.6))
    except Exception:
        logger.warning("logo skipped: unrecognized or invalid image bytes")
        paragraph._element.getparent().remove(paragraph._element)


def _add_header(
    document,
    *,
    brand_name: str | None,
    document_label: str | None,
    date_text: str | None,
    font: str | None,
    accent: RGBColor | None,
    logo_bytes: bytes | None = None,
) -> None:
    """Верхний блок (letterhead): логотип, бренд, мета `label · date`, линия.

    Рендерится, если задан хоть один из logo_bytes/brand_name/document_label/date_text.
    """
    if not (logo_bytes or brand_name or document_label or date_text):
        return

    if logo_bytes:
        _add_logo(document, logo_bytes)

    if brand_name:
        run = document.add_paragraph().add_run(brand_name)
        run.bold = True
        run.font.size = Pt(15)
        if accent is not None:
            run.font.color.rgb = accent
        if font:
            run.font.name = font

    meta = " · ".join(part for part in (document_label, date_text) if part)
    if meta:
        run = document.add_paragraph().add_run(meta)
        run.font.size = Pt(9)
        run.font.color.rgb = _META_GREY
        if font:
            run.font.name = font

    _add_divider(document, accent if accent is not None else _NEUTRAL_DIVIDER)


def _add_inline(paragraph, text: str) -> None:
    """Добавить текст в параграф, превращая `**...**` в жирные run-ы."""
    last = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > last:
            paragraph.add_run(text[last : m.start()])
        paragraph.add_run(m.group(1)).bold = True
        last = m.end()
    if last < len(text):
        paragraph.add_run(text[last:])


def _style_runs(
    paragraph, *, font: str | None = None, color: RGBColor | None = None
) -> None:
    """Применить бренд-шрифт/цвет к run-ам параграфа (если заданы)."""
    for run in paragraph.runs:
        if font:
            run.font.name = font
        if color is not None:
            run.font.color.rgb = color


def build_docx(
    markdown: str,
    *,
    title: str | None = None,
    identity: dict[str, Any] | None = None,
    brand_name: str | None = None,
    document_label: str | None = None,
    date_text: str | None = None,
    logo_bytes: bytes | None = None,
) -> bytes:
    """Собрать .docx (bytes) из markdown.

    `title` пишется в core-properties документа, но НЕ дублируется в теле —
    markdown обычно уже начинается с `# {title}`.

    `identity` — опциональный brand_identity_json. Применяются только
    `font_family` (шрифт документа) и `accent_color`/`primary_color` (цвет
    заголовков). Пустая/невалидная identity (`None`/`{}`/битый hex) → DOCX
    идентичен прежнему. Логотип не используется.

    `brand_name`/`document_label`/`date_text`/`logo_bytes` — верхний блок
    документа (letterhead) перед телом. `logo_bytes` (готовые PNG/JPEG-байты,
    например из logo_fetcher) вставляются картинкой сверху; нераспознанные —
    молча пропускаются. Если ничего из этого не задано — блок не добавляется и
    DOCX идентичен прежнему (backward compatible).
    """
    identity = identity or {}
    raw_font = identity.get("font_family")
    font = raw_font.strip() if isinstance(raw_font, str) and raw_font.strip() else None
    heading_color = _parse_hex(identity.get("accent_color")) or _parse_hex(
        identity.get("primary_color")
    )

    document = Document()
    if title:
        document.core_properties.title = title
    if font:
        document.styles["Normal"].font.name = font

    _add_header(
        document,
        brand_name=brand_name,
        document_label=document_label,
        date_text=date_text,
        font=font,
        accent=heading_color,
        logo_bytes=logo_bytes,
    )

    for raw in (markdown or "").replace("\r\n", "\n").split("\n"):
        line = raw.strip()
        if not line:
            continue
        heading = _HEADING_RE.match(line)
        if heading:
            level = len(heading.group(1))
            paragraph = document.add_paragraph(style=f"Heading {level}")
            _add_inline(paragraph, heading.group(2).strip())
            _style_runs(paragraph, font=font, color=heading_color)
            continue
        bullet = _BULLET_RE.match(line)
        if bullet:
            paragraph = document.add_paragraph(style="List Bullet")
            _add_inline(paragraph, bullet.group(1).strip())
            _style_runs(paragraph, font=font)
            continue
        paragraph = document.add_paragraph()
        _add_inline(paragraph, line)
        _style_runs(paragraph, font=font)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
